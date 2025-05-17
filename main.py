import sys
from docx import Document

def populate_docx(full_text: str, output_file="output.docx"):
    doc = Document()
    doc.add_heading("Digitized Document", level=1)
    doc.add_paragraph(full_text)
    doc.save(output_file)
    print(f"Document saved as: {output_file}")

if __name__ == "__main__":
    text = sys.stdin.read()
    populate_docx(text)
