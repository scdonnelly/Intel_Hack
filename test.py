from pathlib import Path
import cv2
import numpy as np
import pytesseract
import subprocess
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import openvino as ov
import sys
from docx import Document

app = Flask(__name__)
CORS(app)

THRESHOLD = 0.1


# Download the model from the OpenVINO Model Zoo if not already downloaded
def download_file(
    url: str,
    filename: str = None,
    directory: str = None,
    show_progress: bool = True,
) -> Path:
    from tqdm.notebook import tqdm_notebook
    import requests
    import urllib.parse

    filename = filename or Path(urllib.parse.urlparse(url).path).name
    chunk_size = 16384  # make chunks bigger so that not too many updates are triggered for Jupyter front-end

    # check if the filename is a valid file name
    filename = Path(filename)
    if len(filename.parts) > 1:
        raise ValueError(
            "`filename` should refer to the name of the file, excluding the directory. "
            "Use the `directory` parameter to specify a target directory for the downloaded file."
        )

    # check if the directory is a valid directory
    filepath = Path(directory) / filename if directory is not None else filename

    # check if the file already exists
    if filepath.exists():
        return filepath.resolve()

    # create the directory if it does not exist, and add the directory to the filename
    if directory is not None:
        Path(directory).mkdir(parents=True, exist_ok=True)

    # download the file
    try:
        response = requests.get(
            url=url, headers={"User-agent": "Mozilla/5.0"}, stream=True
        )
        response.raise_for_status()
    except requests.exceptions.HTTPError as error:  # For error associated with not-200 codes. Will output something like: "404 Client Error: Not Found for url: {url}"
        raise Exception(error) from None
    except requests.exceptions.Timeout:
        raise Exception(
            "Connection timed out. If you access the internet through a proxy server, please "
            "make sure the proxy is set in the shell from where you launched Jupyter."
        ) from None
    except requests.exceptions.RequestException as error:
        raise Exception(f"File downloading failed with error: {error}") from None

    response.close()

    return filepath.resolve()


def detect_text_regions(image):
    # Check if the model directory exists, if not create it
    base_model_dir = Path("./model").expanduser()

    # Define the model name and paths
    model_name = "horizontal-text-detection-0001"
    model_xml_name = f"{model_name}.xml"
    model_bin_name = f"{model_name}.bin"

    # Set the model paths
    model_xml_path = base_model_dir / model_xml_name
    model_bin_path = base_model_dir / model_bin_name

    # Check if the model files exist, if not download them
    if not model_xml_path.exists():
        model_xml_url = "https://storage.openvinotoolkit.org/repositories/open_model_zoo/2022.3/models_bin/1/horizontal-text-detection-0001/FP32/horizontal-text-detection-0001.xml"
        model_bin_url = "https://storage.openvinotoolkit.org/repositories/open_model_zoo/2022.3/models_bin/1/horizontal-text-detection-0001/FP32/horizontal-text-detection-0001.bin"

        download_file(model_xml_url, model_xml_name, base_model_dir)
        download_file(model_bin_url, model_bin_name, base_model_dir)
    else:
        print(f"{model_name} already downloaded to {base_model_dir}")

    core = ov.Core()

    # Load the model
    model = core.read_model(model=model_xml_path)
    device = ov.Core().get_available_devices()[0]
    compiled_model = core.compile_model(model=model, device_name=device)

    # Get the input and output layers
    input_layer_ir = compiled_model.input(0)
    output_layer_ir = compiled_model.output("boxes")

    # N,C,H,W = batch size, number of channels, height, width.
    N, C, H, W = input_layer_ir.shape

    # Resize the image to meet network expected input sizes.
    resized_image = cv2.resize(image, (W, H))

    # Reshape to the network input shape.
    input_image = np.expand_dims(resized_image.transpose(2, 0, 1), 0)

    # plt.imshow(cv2.cvtColor(image, cv2.COLOR_BGR2RGB));

    # Create an inference request.
    boxes = compiled_model([input_image])[output_layer_ir]

    # Remove zero only boxes.
    boxes = boxes[~np.all(boxes == 0, axis=1)]

    # Get the number of boxes
    total_boxes = len(boxes)
    print("Total boxes: ", total_boxes)

    # Set the initial min and max values to the first box and i to 1
    i = 1
    min_x = boxes[0][0]
    max_x = boxes[0][2]
    min_y = boxes[0][1]
    max_y = boxes[0][3]

    # Loop through the boxes to find the min and max values
    while i < total_boxes:
        if boxes[i][0] < min_x:
            min_x = boxes[i][0]
        if boxes[i][2] > max_x:
            max_x = boxes[i][2]
        if boxes[i][1] < min_y:
            min_y = boxes[i][1]
        if boxes[i][3] > max_y:
            max_y = boxes[i][3]
        i += 1

    # Fing the original image dimensions
    original_height, original_width = image.shape[:2]

    # Calculate the width and height multipliers
    Width_multiplier = original_width / W
    Height_multiplier = original_height / H

    # Ensure bounding box coordinates are integers
    min_x = int(min_x * Width_multiplier - 1)
    max_x = int(max_x * Width_multiplier + 1)
    min_y = int(min_y * Height_multiplier - 1)
    max_y = int(max_y * Height_multiplier + 1)

    print(f"min_x: {min_x}, max_x: {max_x}, min_y: {min_y}, max_y: {max_y}")

    # Validate bounding box
    if min_x >= max_x or min_y >= max_y:
        print("Error: Invalid bounding box coordinates.")
        print(f"min_x: {min_x}, max_x: {max_x}, min_y: {min_y}, max_y: {max_y}")
        exit()

    # Crop the image using integer coordinates
    image = image[min_y:max_y, min_x:max_x]

    # Draw bounding boxes on the original image
    for box in boxes:
        x_min, y_min, x_max, y_max = map(
            int, box[:4]
        )  # Convert box coordinates to integers
        cv2.rectangle(
            image, (x_min, y_min), (x_max, y_max), (0, 255, 0), 2
        )  # Green box with thickness 2

    return boxes


def run_ocr_on_boxes(image, boxes):
    results = []
    print(boxes)
    for i, (x1, y1, x2, y2, p) in enumerate(boxes):
        roi = image[int(y1) : int(y2), int(x1) : int(x2)]
        if roi.size == 0:
            continue  # skip invalid crops
        text = pytesseract.image_to_string(roi, config="--psm 6")
        if text.strip():
            results.append(text.strip())
    return "\n".join(results)


def populate_docx(full_text: str, output_file="output.docx"):
    doc = Document()
    doc.add_heading("Digitized Document", level=1)
    doc.add_paragraph(full_text)
    doc.save(output_file)
    print(f"Document saved as: {output_file}")
    return output_file


@app.route("/")
def hello_world():
    return "<p>Hello, World!</p>"


@app.route("/process-image", methods=["POST"])
def process_image():
    image = request.files["image"]
    doc_type = request.form.get("values")

    if image.filename == "":
        return jsonify({"error": "Image file is empty"}), 400

    try:
        img_bytes = image.read()
        file_bytes = np.frombuffer(img_bytes, np.uint8)
        image = cv2.imdecode(file_bytes, cv2.IMREAD_COLOR)
        if image is None:
            return jsonify({"error": "Image decoding failed"}), 400
    except Exception as e:
        return jsonify({"error": f"Failed to decode image: {str(e)}"}), 500

    boxes = detect_text_regions(image)
    print(f"Detected {len(boxes)} text regions.")

    extracted_text = run_ocr_on_boxes(image, boxes)
    print("Extracted text:")
    print(extracted_text)

    return_doc = populate_docx(extracted_text)

    # Return text for debug
    # Return the DOCX file as an attachment
    return send_file(
        return_doc,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=f"extracted_text_{doc_type}.docx",
    )
